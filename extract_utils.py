
import requests
import sys
import fitz
from asn1crypto import cms
from ocr_extract import extract_text_from_scanned_pdf
import tempfile
import os
import pyarrow.parquet as pq
import pyarrow.compute as pc
from tqdm import tqdm
import docx
import io
import magic

# need to remember cookies for extracting the document
session = requests.Session()
session.headers.update({
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
    'Accept': 'application/json, text/plain, */*',
    'Referer': 'https://e-licitatie.ro/pub/notices/contract-notices/list/0/0'
})

def get_file_type(file_bytes):
    if not file_bytes:
        return None
    mime = magic.from_buffer(file_bytes, mime=True)
    return mime

# get the list of docs
def get_Cnotice_docs(cnotice_id):
    headers = {'Content-Type': 'application/json;charset=UTF-8',
               'Referer': 'https://e-licitatie.ro/pub/notices/contract-notices/list/0/0'}
    url = f'https://www.e-licitatie.ro/api-pub/NoticeCommon/GetDfNoticeSectionFiles/?initNoticeId={cnotice_id}&sysNoticeTypeId=2'
    r = session.get(url, timeout=10)
    return r.json()

# get the specific document
def get_document(url):
    url = f'https://www.e-licitatie.ro/{url}'
    #print(url)
    r = session.get(url, timeout=30, stream=True)
    #print(r)
    return r.content

# go through the list of documents and extract the relevant one
def extract_specifications_url(response):
    doc_list = response.get('dfNoticeDocs', [])
    hints = ['caiet', 'sarcini', 'Caiet', 'Sarcini', 'CAIET', 'SARCINI', 'CS', 'cs', 'C.S.', 'c.s.']
    avoid_hints = ['anexa', 'Anexa', 'ANEXA']
    for doc in doc_list:
        if any(hint in doc.get('noticeDocumentName', '') for hint in hints) and not any(hint in doc.get('noticeDocumentName', '') for hint in avoid_hints):
            specs = get_document(doc.get('noticeDocumentUrl', ''))
            return specs
    return ''

# extract the pdf from digitally signed file
def extract_pdf_from_p7s(p7s_bytes):
    try:
        content_info = cms.ContentInfo.load(p7s_bytes)
        compressed_data = content_info['content']['encap_content_info']['content'].native

        if isinstance(compressed_data, bytes):
            return compressed_data
        else:
            return None
    except Exception as e:
        print(f"error extracting pdf from p7s: {e}")

# extract the raw text from a pdf
def extract_pdf_text(pdf_file):
    try:
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        text_content = ""
        for page in doc:
             text_content += page.get_text()
        doc.close()
        return text_content
    except Exception as e:
         print(f"error extracting raw text from pdf: {e}")

def extract_docx_text(file):
    with io.BytesIO(file) as docx_stream:
        doc = docx.Document(docx_stream)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        fullText.append(" | ".join(row_data))
    return '\n'.join(fullText)

def extract_text_file(doc):
    mime_type = get_file_type(doc)

    if "pkcs7" in mime_type or "octet-stream" in mime_type:
        doc = extract_pdf_from_p7s(doc)
        mime_type = get_file_type(doc)
    
    if mime_type == 'application/pdf':
        # extract text from pdf
        text_final = extract_pdf_text(doc)
        # scanned document
        if len(text_final) < 2000:
            # have to write to a temporary file for paddleocr
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(doc)
                tmp_path = tmp.name
            try:
                text_final = extract_text_from_scanned_pdf(tmp_path)
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
    # docx
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        text_final = extract_docx_text(doc)
    else:
        raise Exception(f"Unsupported file type: {mime_type}")
    os.makedirs("caiete_text", exist_ok = True)
    file_path = os.path.join("caiete_text", f"caiet_sarcini_{cnid}.txt")
    with open(file_path, "w", encoding="utf-8") as f:
                f.write(text_final)

def extract_text_cnid(cnid):
    rez = get_Cnotice_docs(cnid)
    #print(rez)
    doc = extract_specifications_url(rez)
    return extract_text_file(doc)



def process_ca_dataset():
    table = pq.read_table('seap_dataset/contract_awards/', columns=['cNoticeId'])
    # filter notices that dont have cnoticeid and convert to a list
    filtered_table = table.filter(pc.field("cNoticeId").is_valid())
    records = filtered_table.to_pylist()

    pbar = tqdm(records, desc="Cnotice_ids", position=0, leave=False)
    for record in pbar:
        try:
            if os.path.exists(os.path.join("caiete_text", f"caiet_sarcini_{record['cNoticeId']}.txt")):
                continue
            extract_text_cnid(record['cNoticeId'])
        except Exception as e:
            tqdm.write(f"Exception processing {record['cNoticeId']}, exception {e}")

if __name__ == "__main__":
    #extract_text_cnid(sys.argv[1])
    process_ca_dataset()
